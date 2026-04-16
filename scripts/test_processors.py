import sys
import os
import docx

# Add project root to sys.path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.core.cv_processor import CVProcessor
from app.core.job_processor import JobProcessor

def test_processors():
    print("Testing JobProcessor...")
    jp = JobProcessor()
    raw_text = "   Software  Engineer    with python skills.   "
    processed = jp.process(raw_text)
    print(f"Raw: '{raw_text}'")
    print(f"Propcessed: '{processed['text']}'")
    
    if processed['text'] == "Software Engineer with python skills.":
        print("SUCCESS: JobProcessor cleaned text correctly.")
    else:
        print("FAILURE: JobProcessor cleaning failed.")

    print("\nTesting CVProcessor (DOCX)...")
    cp = CVProcessor()
    
    # Create dummy DOCX
    docx_path = "test_cv.docx"
    doc = docx.Document()
    doc.add_paragraph("Curriculum Vitae")
    doc.add_paragraph("Name: John Doe")
    doc.add_paragraph("Skills: Python, AI")
    doc.save(docx_path)
    
    try:
        result = cp.process(docx_path)
        print(f"Extracted Text:\n{result['text']}")
        if "John Doe" in result['text'] and "Python" in result['text']:
            print("SUCCESS: DOCX extraction working.")
        else:
            print("FAILURE: Text not extracted correctly.")
    except Exception as e:
        print(f"FAILURE: Error processing DOCX: {e}")
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)

if __name__ == "__main__":
    test_processors()
